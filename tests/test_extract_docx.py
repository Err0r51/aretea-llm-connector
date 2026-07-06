"""AC2: extract_docx returns paragraph + table-cell text in document order (real .docx)."""

from __future__ import annotations

import io

from aretea_drive_mcp.gdrive.extract import extract_docx
from docx import Document

CEIL = 314_572_800


def _docx_bytes() -> bytes:
    """A real .docx: paragraph → table → paragraph, so ordering is observable."""
    doc = Document()
    doc.add_paragraph("INTRO paragraph")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "CELL-LEFT"
    table.cell(0, 1).text = "CELL-RIGHT"
    doc.add_paragraph("OUTRO paragraph")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_docx_text_extracted_in_document_order() -> None:
    text = extract_docx(_docx_bytes(), max_uncompressed_bytes=CEIL)
    assert "INTRO paragraph" in text
    assert "CELL-LEFT" in text
    assert "CELL-RIGHT" in text
    assert "OUTRO paragraph" in text
    # The body walk preserves interleaving: paragraph, then table, then paragraph.
    assert text.index("INTRO") < text.index("CELL-LEFT") < text.index("OUTRO")
